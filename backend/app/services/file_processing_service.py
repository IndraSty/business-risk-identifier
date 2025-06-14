import io
import base64
from typing import Union
import logging

# Required dependencies: pip install pypdf python-docx
try:
    import pypdf
    from docx import Document

    FILE_PROCESSING_AVAILABLE = True
except ImportError as e:
    logging.error("Missing dependencies. Install with: pip install pypdf python-docx")
    FILE_PROCESSING_AVAILABLE = False
    raise e

logger = logging.getLogger(__name__)


class FileProcessorService:
    """Service for processing uploaded files and extracting text content"""

    @staticmethod
    def extract_text_from_base64(
        file_data: str, file_type: str, filename: str = None
    ) -> str:
        """Extract text from base64 encoded file data"""
        if not FILE_PROCESSING_AVAILABLE:
            raise RuntimeError(
                "File processing dependencies not installed. Run: pip install PyPDF2 python-docx"
            )

        try:
            # Decode base64 to bytes
            file_bytes = base64.b64decode(file_data)

            logger.info(
                f"Processing {file_type.upper()} file: {filename or 'unnamed'} ({len(file_bytes)} bytes)"
            )

            if file_type.lower() == "pdf":
                return FileProcessorService._extract_pdf_text(file_bytes)
            elif file_type.lower() == "txt":
                return FileProcessorService._extract_txt_text(file_bytes)
            elif file_type.lower() in ["docx", "doc"]:
                return FileProcessorService._extract_docx_text(file_bytes)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error processing {file_type} file {filename}: {e}")
            raise ValueError(f"Failed to extract text from {file_type} file: {str(e)}")

    @staticmethod
    def _extract_pdf_text(file_bytes: bytes) -> str:
        """Extract text from PDF bytes"""
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = pypdf.PdfReader(pdf_file)

            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(
                        f"Error extracting text from page {page_num + 1}: {e}"
                    )
                    continue

            extracted_text = "\n\n".join(text_content)

            if not extracted_text.strip():
                raise ValueError("No readable text found in PDF")

            logger.info(
                f"Successfully extracted {len(extracted_text)} characters from PDF"
            )
            return extracted_text

        except Exception as e:
            raise ValueError(f"PDF processing error: {str(e)}")

    @staticmethod
    def _extract_txt_text(file_bytes: bytes) -> str:
        """Extract text from TXT bytes"""
        try:
            # Try different encodings
            encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

            for encoding in encodings:
                try:
                    text = file_bytes.decode(encoding)
                    logger.info(f"Successfully decoded TXT with {encoding} encoding")
                    return text
                except UnicodeDecodeError:
                    continue

            raise ValueError("Unable to decode text file with supported encodings")

        except Exception as e:
            raise ValueError(f"TXT processing error: {str(e)}")

    @staticmethod
    def _extract_docx_text(file_bytes: bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)

            text_content = []

            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract table text
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text_content.append("\n--- Table ---\n" + "\n".join(table_text))

            extracted_text = "\n\n".join(text_content)

            if not extracted_text.strip():
                raise ValueError("No readable text found in DOCX")

            logger.info(
                f"Successfully extracted {len(extracted_text)} characters from DOCX"
            )
            return extracted_text

        except Exception as e:
            raise ValueError(f"DOCX processing error: {str(e)}")

    @staticmethod
    def validate_file_size(file_data: str, max_size_mb: int = 10) -> bool:
        """Validate file size from base64 data"""
        try:
            # Calculate approximate file size
            file_size_bytes = len(file_data) * 3 / 4  # Base64 overhead
            file_size_mb = file_size_bytes / (1024 * 1024)

            logger.info(f"File size: {file_size_mb:.2f} MB (limit: {max_size_mb} MB)")
            return file_size_mb <= max_size_mb

        except Exception:
            return False


file_processing_service = FileProcessorService()
